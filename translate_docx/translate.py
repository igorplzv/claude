import docx
from deep_translator import GoogleTranslator
from tqdm import tqdm
import os

def translate_docx(input_path, output_path, source_lang, target_lang):
    # Загрузка документа
    doc = docx.Document(input_path)
    translator = GoogleTranslator(source=source_lang, target=target_lang)

    def translate_text(text):
        try:
            return translator.translate(text)
        except Exception as e:
            print(f"Translation error: {e}")
            return text

    # Сбор всех элементов для перевода
    elements = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            elements.append(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        elements.append(paragraph)
    for section in doc.sections:
        for header in [section.header, section.footer]:
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    elements.append(paragraph)

    # Перевод с прогресс-баром
    for element in tqdm(elements, desc="Translating", unit="element"):
        if element.text.strip():
            translated = translate_text(element.text)
            # Сохраняем форматирование первого run
            if element.runs:
                element.runs[0].text = translated
                for run in element.runs[1:]:
                    run.text = ''
            else:
                element.add_run(translated)

    # Сохранение документа
    doc.save(output_path)
    print(f"\nTranslation complete! File saved as {output_path}")

if __name__ == "__main__":
    # Запрос пути к исходному файлу
    input_path = input("Введите путь к исходному файлу DOCX: ").strip()
    if not os.path.exists(input_path):
        print("Файл не найден!")
        exit()

    # Выбор направления перевода
    print("Выберите направление перевода:")
    print("1. С английского на русский")
    print("2. С русского на английский")
    choice = input("Введите номер (1 или 2): ").strip()

    if choice == "1":
        source_lang, target_lang = "en", "ru"
    elif choice == "2":
        source_lang, target_lang = "ru", "en"
    else:
        print("Некорректный выбор!")
        exit()

    # Запрос пути для сохранения переведенного файла
    output_path = input("Введите путь для сохранения переведенного файла (по умолчанию: translated.docx): ").strip()
    if not output_path:
        output_path = "translated.docx"

    # Запуск перевода
    translate_docx(input_path, output_path, source_lang, target_lang)