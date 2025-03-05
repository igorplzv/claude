from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_page_borders(section):
    # Настройка границ страницы (имитация рамки)
    page_border = OxmlElement('w:pgBorders')
    page_border.set(qn('w:rsidR'), '00AB46C5')
    page_border.set(qn('w:rsidRPr'), '00AB46C5')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), 'auto')
        page_border.append(border)
    
    section._sectPr.append(page_border)

def add_custom_paragraph(doc, text, bold=False, alignment='left'):
    p = doc.add_paragraph()
    if alignment == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def create_full_tech_card(output_file="Технологическая_карта_ЕСТД.docx"):
    doc = Document()
    
    # Настройка страницы
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.0)
    
    set_page_borders(section)  # Добавляем границы

    # Шапка документа
    add_custom_paragraph(doc, "ООО «МедТехПроизводство»", bold=True, alignment='center')
    add_custom_paragraph(doc, "ТЕХНОЛОГИЧЕСКАЯ КАРТА", bold=True, alignment='center')
    add_custom_paragraph(doc, "Лист 1 из 5", alignment='center')
    
    # Основные реквизиты
    doc.add_paragraph().add_run("Наименование изделия: Ножка тазобедренного протеза").bold = True
    doc.add_paragraph().add_run("Материал: Ti6Al4V (ВТ6)")
    doc.add_paragraph().add_run("Обозначение: ТК-ПТБ-001-2023")
    doc.add_paragraph().add_run("Статус документа: Утверждена").bold = True
    
    # Технические требования
    add_custom_paragraph(doc, "ТЕХНИЧЕСКИЕ ТРЕБОВАНИЯ:", bold=True)
    tech_reqs = [
        ("1. Геометрические параметры:", [
            "Соответствие КД с допусками ±0,1 мм",
            "Отклонение формы не более 0,05 мм",
            "Отклонение расположения не более 0,1 мм"
        ]),
        ("2. Шероховатость поверхности:", [
            "Ra ≤ 0,8 мкм для функциональных поверхностей",
            "Ra ≤ 1,6 мкм для остальных поверхностей"
        ]),
        ("3. Механические свойства:", [
            "Предел прочности при растяжении: ≥ 900 МПа",
            "Предел текучести: ≥ 800 МПа",
            "Относительное удлинение: ≥ 10 %"
        ])
    ]
    
    for title, items in tech_reqs:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    # Таблица операций
    add_custom_paragraph(doc, "ТЕХНОЛОГИЧЕСКИЙ ПРОЦЕСС:", bold=True)
    
    operations = [
        ("005", "4101", "Входной контроль платформ", 
         "Штангенциркуль\nЛазерный анализатор", 
         "ТИ №123", 
         "Проверка геометрии"),
        ("010", "4260", "Подготовка производства", 
         "ПК, ПО Materialise Magics", 
         "ТИ №456", 
         "Оптимизация модели"),
        # Добавьте остальные операции по аналогии
    ]
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Заголовки столбцов
    hdr = table.rows[0].cells
    hdr[0].text = "№ оп."
    hdr[1].text = "Код"
    hdr[2].text = "Наименование операции"
    hdr[3].text = "Оборудование"
    hdr[4].text = "Документ"
    hdr[5].text = "Примечание"
    
    # Заполнение данных
    for op in operations:
        row = table.add_row().cells
        row[0].text = op[0]
        row[1].text = op[1]
        row[2].text = op[2]
        row[3].text = op[3]
        row[4].text = op[4]
        row[5].text = op[5]

    # Раздел контроля качества
    add_custom_paragraph(doc, "КОНТРОЛЬ КАЧЕСТВА:", bold=True)
    quality_control = [
        ("Входной контроль:", [
            "Проверка сертификатов",
            "Контроль геометрии",
            "Проверка документации"
        ]),
        ("Операционный контроль:", [
            "Контроль режимов оборудования",
            "Проверка промежуточных размеров"
        ]),
        ("Приемочный контроль:", [
            "КИМ-измерения",
            "Контроль шероховатости",
            "Механические испытания"
        ])
    ]
    
    for title, items in quality_control:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    # Подписи
    doc.add_page_break()
    add_custom_paragraph(doc, "Разработал:", bold=True)
    add_custom_paragraph(doc, "Инженер-технолог: _________________ /Иванов А.А./")
    add_custom_paragraph(doc, "Проверил: _________________________ /Петров С.И./")
    add_custom_paragraph(doc, "Утвердил: _________________________ /Сидоров В.К./")

    doc.save(output_file)
    print(f"Документ создан: {output_file}")

if __name__ == "__main__":
    create_full_tech_card()