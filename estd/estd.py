from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import time

def add_formatted_paragraph(doc, text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return paragraph

def format_table_cell(cell, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold
    return cell

def add_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_num = paragraph.add_run()
    page_num.add_field('PAGE')

def create_main_stamp(doc):
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Установка ширины столбцов
    widths = [4, 2, 2, 4, 3, 2]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Cm(width)
    
    # Объединение и заполнение ячеек
    cell_org = table.cell(0, 0)
    cell_org.merge(table.cell(0, 2))
    cell_org.text = "НПМК 'МЕДПРОМ'"
    
    cell_doc = table.cell(0, 3)
    cell_doc.merge(table.cell(0, 4))
    cell_doc.text = "ТЕХНОЛОГИЧЕСКАЯ КАРТА"
    
    table.cell(0, 5).text = "О₁"
    table.cell(1, 0).text = "МТК.01.00001"

def create_operation_card(doc, operation):
    add_formatted_paragraph(doc, f"Операция {operation['number']} {operation['code']} {operation['name']}", 
                          bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.allow_autofit = False
    
    # Установка ширины столбцов
    table.columns[0].width = Cm(10)
    table.columns[1].width = Cm(8)
    
    # Заголовки
    header_cells = table.rows[0].cells
    format_table_cell(header_cells[0], "Оборудование", True, WD_ALIGN_PARAGRAPH.CENTER)
    format_table_cell(header_cells[1], "Инструменты и приспособления", True, WD_ALIGN_PARAGRAPH.CENTER)
    
    # Добавление данных
    data_row = table.add_row()
    equip_cell = data_row.cells[0]
    tools_cell = data_row.cells[1]
    
    for item in operation['equipment']:
        equip_cell.add_paragraph(f"• {item}")
    
    for item in operation['tools']:
        tools_cell.add_paragraph(f"• {item}")
    
    # Добавление содержания операции
    doc.add_paragraph()
    add_formatted_paragraph(doc, "Содержание операции:", bold=True)
    for i, step in enumerate(operation['steps'], 1):
        doc.add_paragraph(f"{i}. {step}")
    
    # Добавление контроля
    doc.add_paragraph()
    add_formatted_paragraph(doc, "Контроль:", bold=True)
    for item in operation['control']:
        doc.add_paragraph(f"• {item}")
    
    # Добавление документации
    doc.add_paragraph()
    add_formatted_paragraph(doc, "Документация:", bold=True)
    doc.add_paragraph(operation['docs'])

def create_operations_list():
    operations = [
        {
            'number': '005',
            'code': '4101',
            'name': 'Входной контроль',
            'equipment': [
                'Лазерный анализатор частиц Malvern Panalytical',
                'Вибросито контрольное ВС-2',
                'Весы аналитические AND HR-250AZ',
                'Микроскоп металлографический'
            ],
            'tools': [
                'Набор сит (15-45 мкм)',
                'Контейнеры для проб',
                'Совок пробоотборный'
            ],
            'steps': [
                'Отобрать пробы порошка Ti6Al4V из каждой партии',
                'Проверить гранулометрический состав (фракция 15-45 мкм, не менее 90%)',
                'Проверить насыпную плотность (≥4.41 г/см³)',
                'Проверить текучесть (≤25 с/50г)',
                'Проверить химический состав (Ti - основа, Al 5.5-6.75%, V 3.5-4.5%, Fe ≤0.3%, O ≤0.2%, N ≤0.05%)',
                'Проверить влажность порошка (≤0.1%)',
                'Оформить протокол входного контроля'
            ],
            'control': [
                'Размер частиц: 15-45 мкм',
                'Насыпная плотность: ≥4.41 г/см³',
                'Текучесть: ≤25 с/50г',
                'Химический состав по ГОСТ',
                'Влажность: ≤0.1%'
            ],
            'docs': 'ГОСТ 23849-87'
        },
        {
            'number': '010',
            'code': '4260',
            'name': 'Подготовка производства',
            'equipment': [
                'Рабочая станция CAD HP Z6 G4',
                'Монитор 27" 4K',
                'Принтер лазерный'
            ],
            'tools': [
                'ПО SolidWorks 2024',
                'ПО Materialise Magics',
                'ПО EOS Print 2.0'
            ],
            'steps': [
                'Загрузить 3D-модель в формате STL',
                'Проверить целостность модели и качество поверхностей',
                'Оптимизировать ориентацию детали (угол к платформе ≥45°)',
                'Спроектировать поддерживающие структуры (толщина 0.4-0.6 мм)',
                'Создать слайсы модели (толщина слоя 30 мкм)',
                'Разработать стратегию сканирования',
                'Сгенерировать управляющую программу',
                'Провести виртуальную симуляцию процесса',
                'Сохранить и проверить все файлы'
            ],
            'control': [
                'Проверка целостности STL',
                'Корректность поддержек',
                'Оптимальность ориентации',
                'Качество слайсинга'
            ],
            'docs': 'ТИ 25.89.001'
        },
        {
            'number': '015',
            'code': '5011',
            'name': 'Аддитивное производство',
            'equipment': [
                'SLM-установка EOS M290',
                'Система фильтрации',
                'Система подачи аргона',
                'Чиллер'
            ],
            'tools': [
                'Платформа построения',
                'Скребок для порошка',
                'Щетка специальная',
                'Комплект ключей'
            ],
            'steps': [
                'Подготовить рабочую камеру установки:',
                '  - Очистить камеру от остатков порошка',
                '  - Проверить фильтры системы фильтрации',
                '  - Очистить оптику лазерной системы',
                'Установить и выровнять платформу построения (отклонение ≤0.1 мм)',
                'Загрузить порошок Ti6Al4V (2.5 кг)',
                'Провести продувку камеры аргоном (O₂ ≤0.1%)',
                'Запустить процесс построения:',
                '  - Мощность лазера: 280-300 Вт',
                '  - Скорость сканирования: 800-1000 мм/с',
                '  - Толщина слоя: 30 мкм',
                '  - Температура платформы: 150°C',
                'Контролировать параметры процесса каждые 2 часа',
                'После завершения печати охладить камеру до 40°C',
                'Извлечь платформу с деталью'
            ],
            'control': [
                'Уровень O₂ в камере',
                'Температура платформы',
                'Параметры лазера',
                'Качество наносимого слоя'
            ],
            'docs': 'ТИ 25.89.002'
        },
        {
            'number': '020',
            'code': '0180',
            'name': 'Термическая обработка',
            'equipment': [
                'Вакуумная печь СЭВФ-3.3/11,5-ИО1',
                'Система контроля температуры',
                'Система вакуумирования',
                'Система охлаждения'
            ],
            'tools': [
                'Термопары типа ТПП',
                'Подставки керамические',
                'Защитные перчатки',
                'Пирометр'
            ],
            'steps': [
                'Очистить детали от загрязнений',
                'Установить детали на подставки',
                'Загрузить детали в печь',
                'Создать вакуум (10⁻⁵ торр)',
                'Провести отжиг для снятия напряжений:',
                '  - Нагрев до 750°C со скоростью 100°C/час',
                '  - Выдержка 4 часа',
                '  - Охлаждение с печью до 200°C',
                'Контролировать параметры процесса',
                'Извлечь детали после охлаждения до 60°C',
                'Провести визуальный контроль'
            ],
            'control': [
                'Температура процесса ±5°C',
                'Уровень вакуума 10⁻⁵ торр',
                'Время выдержки ±5 мин',
                'Скорость охлаждения',
                'Отсутствие окисления'
            ],
            'docs': 'ТИ 25.89.003'
        },
        {
            'number': '025',
            'code': '4111',
            'name': 'Механическая обработка',
            'equipment': [
                '5-осевой обрабатывающий центр DMG MORI DMU 50',
                'Система подачи СОЖ',
                'Система удаления стружки',
                'Система ЧПУ Siemens 840D'
            ],
            'tools': [
                'Фрезы твердосплавные VHM Ø4-12 мм',
                'Державки инструментальные HSK-A63',
                'Комплект калибров',
                'Микрометр цифровой 0-25 мм',
                'Штангенциркуль цифровой 0-150 мм'
            ],
            'steps': [
                'Изучить чертеж и технические требования',
                'Установить деталь в приспособление',
                'Произвести привязку инструмента',
                'Обработать поверхности согласно КД:',
                '  - Черновая обработка (припуск 0.5 мм)',
                '  - Получистовая обработка (припуск 0.2 мм)',
                '  - Чистовая обработка',
                'Режимы резания при чистовой обработке:',
                '  - Скорость резания: 60-80 м/мин',
                '  - Подача: 0.1 мм/зуб',
                '  - Глубина резания: 0.2 мм',
                'Контроль размеров',
                'Снять деталь, очистить от стружки'
            ],
            'control': [
                'Точность размеров по КД ±0.1 мм',
                'Шероховатость Ra 3.2',
                'Отсутствие сколов и зарезов',
                'Соответствие геометрии КД'
            ],
            'docs': 'ТИ 25.89.004'
        },
        {
            'number': '030',
            'code': '4115',
            'name': 'Финишная обработка',
            'equipment': [
                'Полировальный станок Struers LaboPol-30',
                'Установка ультразвуковой очистки УЗО-10-1',
                'Сушильный шкаф SNOL 60/300',
                'Система очистки воды'
            ],
            'tools': [
                'Полировальные круги (120-2500 грит)',
                'Алмазные пасты (14/10 - 1/0 мкм)',
                'Материал протирочный безворсовый',
                'Спирт технический',
                'Набор щеток'
                ],
            'steps': [
                'Провести входной контроль детали',
                'Предварительная полировка:',
                '  - Шлифовка кругами 120-320 грит',
                '  - Промывка, контроль поверхности',
                'Промежуточная полировка:',
                '  - Полировка кругами 400-800 грит',
                '  - Промывка, контроль поверхности',
                'Финишная полировка:',
                '  - Полировка кругами 1200-2500 грит',
                '  - Финишная доводка алмазной пастой',
                'Ультразвуковая очистка:',
                '  - Очистка в УЗВ (15 минут)',
                '  - Промывка деионизированной водой',
                'Сушка в сушильном шкафу (60°C, 30 мин)',
                'Контроль качества поверхности'
            ],
            'control': [
                'Шероховатость Ra ≤0.8 мкм',
                'Отсутствие царапин и рисок',
                'Равномерность полировки',
                'Чистота поверхности'
            ],
            'docs': 'ТИ 25.89.005'
        },
        {
            'number': '035',
            'code': '0220',
            'name': 'Контроль',
            'equipment': [
                'КИМ Zeiss CONTURA',
                'Профилометр Hommel-Etamic W5',
                'Микроскоп металлографический Olympus GX51',
                'Твердомер Роквелла ТР-5014',
                'Установка рентгеновского контроля'
            ],
            'tools': [
                'Комплект эталонов шероховатости',
                'Индикаторы часового типа ИЧ-10',
                'Микрометры цифровые',
                'Штангенциркули цифровые',
                'Набор калибров'
            ],
            'steps': [
                'Подготовить измерительное оборудование',
                'Провести калибровку приборов',
                'Проверить геометрические размеры:',
                '  - Измерение на КИМ',
                '  - Контроль критических размеров',
                'Измерить шероховатость поверхностей',
                'Проверить механические свойства:',
                '  - Предел прочности ≥900 МПа',
                '  - Предел текучести ≥800 МПа',
                '  - Относительное удлинение ≥10%',
                'Провести рентгеновский контроль',
                'Оформить протокол контроля',
                'Занести данные в систему контроля качества'
            ],
            'control': [
                'Соответствие размеров КД (±0.1 мм)',
                'Шероховатость Ra ≤0.8 мкм',
                'Механические свойства по ТУ',
                'Отсутствие внутренних дефектов',
                'Качество поверхности'
            ],
            'docs': 'ТИ 25.89.006'
        }
    ]
    return operations

def add_technical_requirements(doc):
    add_formatted_paragraph(doc, "ТЕХНИЧЕСКИЕ ТРЕБОВАНИЯ:", bold=True, size=14)
    requirements = [
        ("1. Геометрические параметры", [
            "- Соответствие КД с допусками ±0,1 мм",
            "- Отклонение формы не более 0,05 мм",
            "- Отклонение расположения не более 0,1 мм"
        ]),
        ("2. Шероховатость поверхности", [
            "- Ra ≤ 0,8 мкм для функциональных поверхностей",
            "- Ra ≤ 1,6 мкм для остальных поверхностей"
        ]),
        ("3. Механические свойства", [
            "- Предел прочности при растяжении: ≥ 900 МПа",
            "- Предел текучести: ≥ 800 МПа",
            "- Относительное удлинение: ≥ 10 %",
            "- Твердость: 30-35 HRC"
        ])
    ]
    
    for title, items in requirements:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

def add_quality_control_section(doc):
    add_formatted_paragraph(doc, "КОНТРОЛЬ КАЧЕСТВА:", bold=True, size=14)
    
    control_items = [
        ("Входной контроль", [
            "Проверка сертификатов на материалы",
            "Контроль геометрических параметров заготовки",
            "Проверка сопроводительной документации"
        ]),
        ("Операционный контроль", [
            "Контроль режимов работы оборудования",
            "Проверка промежуточных размеров",
            "Контроль параметров техпроцесса"
        ]),
        ("Приемочный контроль", [
            "Контроль геометрических параметров",
            "Проверка шероховатости",
            "Контроль механических свойств",
            "Визуальный контроль качества поверхности"
        ])
    ]
    
    for title, items in control_items:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        for item in items:
            doc.add_paragraph(f"• {item}")

def create_document(doc):
    # Создание титульного листа
    create_main_stamp(doc)
    doc.add_paragraph()
    
    # Общая информация
    add_formatted_paragraph(doc, "ТЕХНОЛОГИЧЕСКАЯ КАРТА", bold=True, size=16, 
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    
    # Информация о детали
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    info = [
        ['Наименование изделия:', 'Ножка тазобедренного протеза'],
        ['Материал:', 'Ti6Al4V (ВТ6)'],
        ['Обозначение:', 'МТК.01.00001'],
        ['Статус документа:', 'Технологическая карта']
    ]
    
    for i, (key, value) in enumerate(info):
        cells = table.rows[i].cells
        cells[0].text = key
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Добавление технических требований
    add_technical_requirements(doc)
    doc.add_paragraph()
    
    # Создание карт операций
    operations = create_operations_list()
    for operation in operations:
        doc.add_page_break()
        create_operation_card(doc, operation)
    
    # Добавление раздела контроля качества
    doc.add_page_break()
    add_quality_control_section(doc)
    
    # Добавление подписей
    doc.add_page_break()
    add_formatted_paragraph(doc, "СОГЛАСОВАНО:", bold=True)
    signatures = [
        "Разработал: _____________ /_________/    Дата: ___.___.2024",
        "Проверил: _____________ /_________/      Дата: ___.___.2024",
        "Нормоконтроль: _____________ /_________/ Дата: ___.___.2024",
        "Утвердил: _____________ /_________/      Дата: ___.___.2024"
    ]
    
    for sign in signatures:
        doc.add_paragraph(sign)

def create_tech_card(output_filename='Технологическая_карта_ЕСТД.docx'):
    try:
        doc = Document()
        
        # Настройка параметров страницы
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        
        # Добавление номеров страниц
        # add_page_number(doc)
        
        # Создание документа
        create_document(doc)
        
        # Сохранение с проверкой
        try:
            doc.save(output_filename)
            print(f"Документ успешно создан: {output_filename}")
        except Exception as e:
            print(f"Ошибка при сохранении документа: {e}")
            alt_filename = f"Технологическая_карта_ЕСТД_{int(time.time())}.docx"
            doc.save(alt_filename)
            print(f"Документ сохранен с альтернативным именем: {alt_filename}")
            
    except Exception as e:
        print(f"Ошибка при создании документа: {e}")
        raise

if __name__ == "__main__":
    try:
        create_tech_card()
    except Exception as e:
        print(f"Критическая ошибка: {e}")